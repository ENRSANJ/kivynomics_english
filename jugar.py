import numpy as np
from kivy.app import App
from kivy.lang import Builder
from kivy.properties import StringProperty, NumericProperty
from docx import Document
from kivy.uix.gridlayout import GridLayout
from kivy.uix.popup import Popup
from kivy.uix.scrollview import ScrollView
from kivymd.uix.dialog import MDDialog
from bimatrix_games import MensajeDeError
from main import VentanaLayout, MasInfoVentana, WrappedLabel

Builder.load_file('jugar.kv')


class SelectPlayer(MDDialog):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.referencia = App.get_running_app().root.ids.jugarid

    def selection1(self):
        self.referencia.player2 = NPCNash('images/johnnash.jpg',
                                          ['''"[...]Pero, después de una vida de búsqueda me digo, ¿Qué es la lógica?\
 ¿Quién decide la razón?"''',
                                           '''"Caballeros, debo recordarles que, mis probabilidades de éxito, aumentan\
 en cada nuevo intento"''',
                                           '"No, no creo en la suerte, pero sí en asignar valor a las cosas"'])
        self.dismiss()
        self.referencia.inicio_juego()

    def selection2(self):
        self.referencia.player2 = NPCMurray('images/rothbard.jpg',
                                            ['''"El Estado obtiene su renta mediante el uso de la compulsión, es decir,\
 la amenaza de la cárcel y la bayoneta"''',
                                             '''"Si los humanos son tan malos, ¿cómo podemos esperar que un gobierno\
 coercitivo, compuesto por humanos, mejore la situación?"''',
                                             '''"La contribución es, pura y simplemente, un robo, un robo a grande\
 y colosal escala"'''])
        self.dismiss()
        self.referencia.inicio_juego()

    def selection3(self):
        self.referencia.player2 = NPCMarx('images/karlmarx.jpg',
                                          ['"¡Proletarios de todos los países, uníos!"',
                                           '''"La historia de todas las sociedades hasta el día de hoy es historia\
 de luchas de clases"''',
                                           '"¡De cada cual según sus capacidades, a cada cual según sus necesidades!"'])
        self.dismiss()
        self.referencia.inicio_juego()


class ConfirmationScreen(MDDialog):
    pass


# Para mostrar los resultados en pantalla
class Popup_resultado(Popup):
    pass


class ResetGame(MDDialog):

    @staticmethod
    def exportar_resultados(resultados):
        # Método para exportar los resultados a Word
        # Create a new Word document
        document = Document()

        # Añadimos nuestros datos al documento
        for key, value in resultados.items():
            document.add_paragraph('STAGE: ' + f'{key}')
            document.add_paragraph('MODEL: ' + f'{value[0]}')
            document.add_paragraph('Market demand function: ' + f'p(X) = {value[1]} - {value[2]}X')
            document.add_paragraph('Total Costs of firm 1: ' + f'{value[3]}x\u2081')
            document.add_paragraph('Total Costs of firm 2: ' + f'{value[3]}x\u2082')
            document.add_paragraph()
            document.add_paragraph('ANSWERS: ')
            document.add_paragraph('Your firm: ')
            document.add_paragraph('Production: ' + f'{value[6]} uds')
            document.add_paragraph('Profit: ' + f'{value[7]} €')
            document.add_paragraph()
            document.add_paragraph("NPC's firm: ")
            document.add_paragraph('Production: ' + f'{value[8]} u')
            document.add_paragraph('Profit: ' + f'{value[9]} €')
            document.add_paragraph()
            document.add_paragraph('Market data: ')
            document.add_paragraph('Price: ' + f'{value[4]}')
            document.add_paragraph('Quantity: ' + f'{value[5]}')
            # Siguiente página
            [document.add_paragraph() for i in range(8)]

        document.add_paragraph('GLOBAL RESULTS:')
        document.add_paragraph('Your total profit: ' + f'{resultados[1][7] + resultados[2][7] + resultados[3][7]} €')
        document.add_paragraph("NPC's total profit: " + f'{resultados[1][9] + resultados[2][9] + resultados[3][9]} €')

        # Save the Word document
        document.save('vsNPC_output.docx')

        alerta = MensajeDeError("The document 'vsNPC_output.docx' has been succesfully created")
        alerta.open()

        # Creamos el output a mostrar en pantalla
        scroll = ScrollView()
        popup = Popup_resultado(content=scroll)
        grid = GridLayout(cols=3, size_hint=(1, None), padding='5dp')
        scroll.add_widget(grid)
        grid.bind(minimum_height=grid.setter('height'))

        for key, value in resultados.items():
            text_list = [
                '[b]STAGE: ' + f'{key}[/b]',
                '\n[u]MODEL:[/u] ' + f'{value[0]}',
                '\n[u]Data[/u]',
                f'p(X) = {value[1]} - {value[2]}X',
                'TC[sub]1[/sub]: ' + f'{value[3]}x[sub]1[/sub]',
                'TC[sub]2[/sub]: ' + f'{value[3]}x[sub]2[/sub]',
                '\n[u]ANSWERS:[/u] ',
                '\n[u]Your firm:[/u] ',
                'Production: ' + f'{value[6]} uds',
                'Profit: ' + f'{value[7]} €',
                '\n[u]NPC Firm:[/u] ',
                'Production: ' + f'{value[8]} uds',
                'Profit: ' + f'{value[9]} €',
                '\n[u]Market data:[/u] ',
                'Price: ' + f'{value[4]} €',
                'Quantity: ' + f'{value[5]} uds',
            ]

            grid2 = GridLayout(cols=1, size_hint=(1, None), padding='5dp')
            grid2.bind(minimum_height=grid2.setter('height'))
            for text in text_list:
                grid2.add_widget(
                    WrappedLabel(text=text, color=(0, 0, 0, 1), size_hint=(1, None), markup=True))
            grid.add_widget(grid2)

        '''text_list2 = [
        'GLOBAL RESULTS:',
        'Your total profit: ' + f'{historial[1][7] + historial[2][7] + historial[3][7]} €',
        ' total NPC: ' + f'{historial[1][9] + historial[2][9] + historial[3][9]} €'
        ]
        for text in text_list2:
            grid.add_widget(
                WrappedLabel(text=text, color=(0, 0, 0, 1), font_size=grid.height * 0.022, size_hint=(1, None)))'''

        popup.open()


class JugarVentana(VentanaLayout):
    player2 = None
    prod1 = 0
    prod2 = 0
    precio = 0
    beneficio1 = 0
    beneficio2 = 0
    cantidad_total = 0
    evento = 0
    historial = {}
    model = 0
    precio2 = 0
    modelos = {1: 'COURNOT', 2: 'STACKELBERG (NPC=leader)', 3: 'STACKELBERG (NPC=follower)', 4: 'BERTRAND'}

    # Asignamos valores aleatorios al inicio del juego
    a = 0
    b = 0
    c = 0

    # Valores a mostrar en pantalla
    comentario = StringProperty('')
    modelo = StringProperty('')
    narrativa = StringProperty('')
    demanda_mercado = StringProperty(f'{a} - {b}x')
    costes_totales = StringProperty(f'{c}x')
    imagen = StringProperty('')
    rival = None
    stage = NumericProperty(0)
    stage1 = StringProperty('')
    stage2 = StringProperty('')
    stage3 = StringProperty('')

    def on_enter(self, *args):
        if self.stage == 0:
            selector = SelectPlayer()
            selector.open()
        else:
            pass

    @staticmethod
    def reiniciar():
        dialog = ConfirmationScreen()
        dialog.open()

    def inicio_juego(self):
        # Asignamos valores aleatorios al inicio del juego
        self.a = np.random.randint(250, 500)
        self.b = np.random.randint(2, 25)
        self.c = np.random.randint(21, 40)

        self.imagen = str(self.player2.img)
        self.comentario = self.player2.frases[np.random.randint(0, 3)]
        self.update_screen()
        self.selecciona_modelo()
        self.stage += 1

    def update_screen(self):
        self.demanda_mercado = f'{self.a} - {self.b}x'
        self.costes_totales = f'{self.c}x'

    def selecciona_modelo(self):
        self.model = np.random.randint(1, 5)

        # Cournot
        if self.model == 1:
            self.modelo = 'COURNOT MODEL'
            self.prod2 = np.around(self.player2.cournot(self.a, self.b, self.c), 3)

        # NPC Stackelberg líder
        elif self.model == 2:
            self.prod2 = round(self.player2.stackelberg1(self.a, self.b, self.c), 3)
            self.modelo = f'STACKELBERG MODEL. YOU ARE THE FOLLOWER. {self.player2.name} produced {self.prod2} u.'

        # NPC Stackelberg seguidor
        elif self.model == 3:
            self.modelo = 'STACKELBERG MODEL. YOU ARE THE LEADER'

        # Bertrand
        else:
            self.modelo = 'BERTRAND MODEL'
            self.precio2 = self.player2.bertrand(self.c)

    def reset_parametros(self):
        self.historial = {}
        self.stage = 0
        self.ids.respuesta.text = ''
        self.stage1 = ''
        self.stage2 = ''
        self.stage3 = ''
        self.imagen = ''
        self.comentario = ''
        self.modelo = ''
        self.narrativa = ''
        self.demanda_mercado = ''
        self.costes_totales = ''
        NPCMurray.tax = 0

    def evento_aleatorio(self):
        # Selección aleatoria del evento y la cuantía
        self.evento = np.random.randint(1, 5)
        cuantia = np.random.randint(3, 10)

        # Si el evento afecta a la demanda, modificamos la demanda (y viceversa)
        if self.evento == 1:
            self.narrativa = '''Recently, the results of a study conducted by the University of Massachusetts were \
published, highlighting the benefits of using your product. This has led to an increase in market demand. Changes have \
been recorded in the demand function.'''
            self.a = self.a + cuantia

        elif self.evento == 2:
            self.narrativa = '''As a result of a global pandemic, the number of people who can access and use your \
product has significantly reduced. Changes have been introduced in the market demand function.'''
            self.a = self.a - cuantia

        elif self.evento == 3:
            self.c = self.c + cuantia
            NPCMurray.tax = NPCMurray.tax - cuantia
            self.narrativa = f'''The Government introduced a tax of {cuantia} units of currency (u.m.) on production. \
The resulting increase in costs has been incorporated into your cost function.'''

        elif self.evento == 4:
            self.c = self.c - cuantia
            NPCMurray.tax = NPCMurray.tax + cuantia
            self.narrativa = f'''The Government introduced a subsidy of {cuantia} units of currency (u.m.) on production. \
The resulting increase in costs has been incorporated into your cost function.'''

        self.update_screen()

    def confirma(self):

        # Comprobamos si el usuario introdujo un valor admitido
        try:
            respuesta = np.around(float(self.ids.respuesta.text), 3)
            if respuesta == 0:
                raise ValueError
        except ValueError:
            alerta = MensajeDeError('Enter valid parameters')
            alerta.open()
            return

        # Si es modelo de Bertrand, la variable de decisión cambia
        if self.model == 4:
            self.precio2 = self.precio2

            if self.precio2 > respuesta:
                self.precio = respuesta
                self.prod1 = np.around((self.a - self.precio) / self.b, 3)
                self.prod2 = 0

            elif self.precio2 < respuesta:
                self.precio = self.precio2
                self.prod1 = 0
                self.prod2 = np.around((self.a - self.precio) / self.b, 3)

            # Por defecto, a igual precio suponemos los clientes se distribuyen equitativamente
            else:
                self.precio = respuesta
                self.prod1 = np.around((self.a - self.precio) / (2 * self.b), 3)
                self.prod2 = np.around((self.a - self.precio) / (2 * self.b), 3)

        # Si el NPC es seguidor, debe conocer nuestra producción
        elif self.model == 3:
            self.prod1 = respuesta
            self.prod2 = np.around(self.player2.stackelberg2(self.a, self.b, self.c, respuesta), 3)
            self.precio = np.around(self.a - (self.b * (self.prod1 + self.prod2)), 3)

        else:
            self.prod1 = respuesta
            self.precio = np.around(self.a - (self.b * (self.prod1 + self.prod2)), 3)

        # Si el precio resultante en el mercado es negativo, se establece un precio igual a 0.001
        if self.precio <= 0:
            self.precio = .001

        self.cantidad_total = np.around(self.prod1 + self.prod2, 3)
        self.beneficio1 = np.around((self.precio - self.c) * self.prod1, 3)
        self.beneficio2 = np.around((self.precio - self.c) * self.prod2, 3)

        # Rellenamos el historial
        self.historial[self.stage] = [self.modelos.get(self.model, self.model), self.a, self.b, self.c, self.precio, self.cantidad_total,
                                      self.prod1, self.beneficio1, self.prod2, self.beneficio2]

        stage_dict = {1: 'stage1', 2: 'stage2', 3: 'stage3'}
        stage_key = stage_dict.get(self.stage)
        setattr(self, stage_key, f'Profit_{self.stage}:  {self.beneficio1}')

        # Pasamos de etapa
        self.ids.respuesta.text = ''
        self.stage += 1

        # Comprobamos si el juego se ha acabado
        if self.stage > 3:
            reseteo = ResetGame()
            reseteo.open()
            return

        else:
            self.evento_aleatorio()
            self.selecciona_modelo()

    def obtenermasinfo(self):
        self.manager.current = 'JugarMasInfoScreen'


class NPCNash:
    def __init__(self, img, frases):
        self.img = img
        self.frases = frases
        self.name = 'John Forbes Nash'

    # Cournot
    @staticmethod
    def cournot(a, b, c):
        return (a - c) / (3 * b)

    # Líder
    @staticmethod
    def stackelberg1(a, b, c):
        return (a - c) / (2 * b)

    # Seguidora
    @staticmethod
    def stackelberg2(a, b, c, x1):
        result = ((a - c) / (2 * b)) - (x1 / 2)
        if result > 0:
            return result
        else:
            return 0

    # Bertrand
    @staticmethod
    def bertrand(c):
        return c


class NPCMurray:
    # Variable que (des)informará al NPCMurray sobre los impuestos del Gobierno
    tax = 0

    def __init__(self, img, frases):
        self.img = img
        self.frases = frases
        self.name = 'Murray Rothbard'

    # Cournot
    @staticmethod
    def cournot(a, b, c):
        return (a - (c + NPCMurray.tax)) / (3 * b)

    # Líder
    @staticmethod
    def stackelberg1(a, b, c):
        return (a - (c + NPCMurray.tax)) / (2 * b)

    # Seguidora
    @staticmethod
    def stackelberg2(a, b, c, x1):
        result = ((a - (c + NPCMurray.tax)) / (2 * b)) - (x1 / 2)
        if result > 0:
            return result
        else:
            return 0

    # Bertrand
    @staticmethod
    def bertrand(c):
        return c + NPCMurray.tax


class NPCMarx:
    def __init__(self, img, frases):
        self.img = img
        self.frases = frases
        self.name = 'Karl Marx'

    # Cournot
    @staticmethod
    def cournot(a, b, c):
        return (a - (1.2 * c)) / (3 * b)

    # Líder
    @staticmethod
    def stackelberg1(a, b, c):
        return (a - (1.2 * c)) / (2 * b)

    # Seguidora
    @staticmethod
    def stackelberg2(a, b, c, x1):
        result = ((a - (1.2*c)) / (2 * b)) - (x1 / 2)
        if result > 0:
            return result
        else:
            return 0

    # Bertrand
    @staticmethod
    def bertrand(c):
        return 1.2 * c


class JugarMasInfoScreen(MasInfoVentana):
    a = '''The game consists of 3 stages in which you, as a company owner, will compete with one of the three available NPCs, who are owners of rival companies. Each of these players has a different predefined behavior, changing the ease with which you can obtain profits in each period:
   
   - John Nash: will always make the most appropriate decisions, taking into account the existing theory in this regard.

   - Murray Rothbard: will generally respond accurately but has a weakness: state intervention. He will ignore established taxes and subsidies granted by the government. He will ignore them when making his production decision; however, his actual costs will be the same as yours, so you can take advantage of this circumstance. Keep in mind that his disapproval of state intervention is permanent, so he will "accumulate" ignorance of past interventions. It is advisable to keep track of this to optimize your decisions. For example:

   TC = 10x
   Stage 1:
   "The Government introduces a tax on production of 2 units of currency (u.m.)."
   Rothbard will produce as if he still had TC = 10x when in reality his costs are higher (12x).
   Stage 2:
   "The Government grants a production subsidy of 4 units of currency (u.m.)."
   Rothbard will produce as if he still had TC = 10x (ignoring both the tax from the previous period and the subsidy from this period) when in reality his actual costs are 8x.

   - Karl Marx: overestimates his costs with the intention of not extracting surplus value from his workers. Specifically, he will produce as if his costs were 20% higher than the actual costs. In this case, the NPC takes into account all events that occur. Furthermore, this 20% is not cumulative, meaning that in each period, he will act as if his costs were 20% higher than yours for that same period:

   TC = 10x
   Karl Marx will produce as if he had TC = 12x.

   It is worth remembering that the confrontation with the computer is not always strictly equitable since, despite having the same total costs at all times, in the Stackelberg model, the order of entry into the market unbalances the scale (significantly favors the leading company).

   Throughout the game, you will have to decide the quantity to produce (Cournot and Stackelberg models) or the price you will set (Bertrand model) considering the NPC's personality, the model to be used in the stage, and, of course, the market demand functions and total costs.

   At the end of the game, you will be able to export the results obtained for each stage to a Word document.

'''
