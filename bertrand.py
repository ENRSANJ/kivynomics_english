import numpy as np
from kivy.lang import Builder
from bimatrix_games import MensajeDeError
from cournot import CournotVentana
from main import MasInfoVentana
from docx import Document
from docx.shared import Inches
from kivy.garden.matplotlib.backend_kivyagg import FigureCanvasKivyAgg
import matplotlib.pyplot as plt
import sympy as sp

Builder.load_file('bertrand.kv')


# Función que resuelve el modelo de Bertrand dados 4 parámetros (f.demanda y costes marginales)
def solve_bertrand(a, b, c, e):

    # Comprobamos si los parámetros son compatibles con una solución razonable del modelo
    if a <= c and a <= e:
        raise ValueError('Enter valid parameters')

    if c < e:
        # La empresa 1 se lleva toda la demanda
        if (c + a) / 2 >= e:
            # Si no puede establecer precio de monopolio, establece
            # un precio marginalmente inferior al Coste Marginal de la empresa 2
            p = e - .001

        else:
            p = (c + a) / 2

        x1_sol = (a - p) / b
        x2_sol = 0

    elif c > e:
        # La empresa 2 se lleva toda la demanda
        if (e + a) / 2 >= c:
            # Si no puede establecer precio de monopolio, establece
            # un precio marginalmente inferior al Coste Marginal de la empresa 2
            p = c - .001

        else:
            p = (e + a) / 2

        x1_sol = 0
        x2_sol = (a - p) / b

    else:
        # Distribuimos la demanda a partes iguales
        p = c
        x1_sol = (a - c) / (2 * b)
        x2_sol = (a - c) / (2 * b)

    x_total = x1_sol + x2_sol
    beneficio1 = (p - c) * x1_sol
    beneficio2 = (p - e) * x2_sol

    return {'p': p, 'q': x_total, 'q1': x1_sol, 'q2': x2_sol, 'profit1': beneficio1, 'profit2': beneficio2}


class BertrandVentana(CournotVentana):

    def obtenermasinfo(self):
        self.manager.current = 'BertrandMasInfoScreen'

    def aleatorio(self):
        self.ids.a.text = str(np.random.randint(11, 30))
        self.ids.b.text = str(np.random.randint(1, 10))
        self.ids.c.text = str(np.random.randint(1, 10))
        self.ids.e.text = str(np.random.randint(1, 10))

    def calcula(self):
        if not self.todo_en_orden():
            return

        # Almacenamos en variables el input del usuario
        a = float(self.ids.a.text)
        b = float(self.ids.b.text)
        c = float(self.ids.c.text)
        e = float(self.ids.e.text)

        # Obtenemos los resultados, redondeados
        try:
            solucion = {key: np.around(value, 3) for key, value in solve_bertrand(a, b, c, e).items()}
        except ValueError:
            alerta = MensajeDeError('Enter valid parameters')
            alerta.open()
            return

        self.precio = 'Market price: ' + str(solucion['p'])
        self.empresa1 = 'FIRM 1'
        self.empresa2 = 'FIRM 2'
        self.prod_1 = 'Optimal production: ' + str(solucion['q1'])
        self.prod_2 = 'Optimal production: ' + str(solucion['q2'])
        self.beneficio1 = 'Profit: ' + str(solucion['profit1'])
        self.beneficio2 = 'Profit: ' + str(solucion['profit2'])

        x1_sol = solucion['q1']
        x2_sol = solucion['q2']
        x_total = solucion['q']
        p = solucion['p']

        # Gráfico
        # Generamos los valores de x
        demanda_max = int(np.ceil((a/b)*1.1))
        x_vals = np.linspace(0, demanda_max)

        # Creamos la función de demanda
        x = sp.symbols('x')
        fl1 = sp.lambdify(x, (a-b*x))

        # Calculamos los valores de p asociados a cada valor de x mediante la función de demanda
        y1 = fl1(x_vals)

        # Dibujamos el gráfico
        fig = plt.figure()
        ax = fig.add_subplot(111)

        # Dibujamos la ecuación
        ax.plot(x_vals, y1, color='b')

        # Dibujamos los costes marginales
        plt.axhline(y=c, color='r', linestyle='--')
        plt.axhline(y=e, color='r', linestyle='--')

        # Dibujamos la producción
        ax.plot(x_total, p, marker='o', markersize=5, c='black', alpha=1)
        ax.annotate(f'Eq ({x_total:.3f}, {p:.3f})', xy=(x_total, p),
                    xycoords='data', xytext=(5, 5), textcoords='offset points')

        # Líneas de puntos entre los ejes y la producción de cada empresa
        ax.plot([x1_sol, x1_sol], [0, p], ':', c='black')
        ax.plot([x2_sol, x2_sol], [0, p], ':', c='black')

        # Editamos las características del gráfico
        plt.title('Bertrand model')
        plt.xlabel('x, x$_{1}$, x$_{2}$')
        plt.ylabel('p')

        ax.set_xlim(0, demanda_max)
        ax.set_ylim(0, int(np.ceil(a*1.1)))
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.legend(handles=[plt.plot([], [], 'b-', label='Demand function')[0],
                  plt.plot([], [], 'r--', label='Marginal Costs')[0]])

        # Cambiamos el grosor de los bordes
        for spine in ['bottom', 'left']:
            ax.spines[spine].set_linewidth(2)

        # Creamos el widget FigureCanvasKivyAgg (vaciamos "graficournot" para eliminar el gráfico anterior)
        canvas = FigureCanvasKivyAgg(figure=fig)
        plt.close("all")
        self.ids.graficournot.clear_widgets()
        self.ids.graficournot.add_widget(canvas)

        # Generamos el output en Word si el usuario lo requiere
        try:
            if self.ids.a_word.active:
                fdemanda = f'p(X) = {a} - {b}X ,    X = x\u2081 + x\u2082'
                ct1 = f'CT(x\u2081) = {c}x\u2081'
                ct2 = f'CT(x\u2082) = {e}x\u2082'

                fig.savefig('bertrand_graph.png')
                crear_bertrand_word(fdemanda, ct1, ct2, self.precio, self.prod_1,
                                    self.beneficio1, self.prod_2, self.beneficio2)
                alerta = MensajeDeError("The Word document 'bertrand_output.docx' has been succesfully created")
                alerta.title = ''
                alerta.open()
        except PermissionError:
            alerta = MensajeDeError("Close all files named 'bertrand_output.docx' or 'bertrand_graph.png'")
            alerta.open()


# Método para exportar los resultados a Word
def crear_bertrand_word(fdemanda, ct1, ct2, precio, prod1, beneficio1, prod2, beneficio2):
    # Create a new Word document
    document = Document()

    # Añadimos nuestros datos al documento
    document.add_paragraph('BERTRAND MODEL')
    document.add_paragraph('Market demand function: ' + fdemanda)
    document.add_paragraph('Total Costs of firm 1: ' + ct1)
    document.add_paragraph('Total Costs of firm 2: ' + ct2)
    document.add_paragraph()
    document.add_picture('bertrand_graph.png', width=Inches(6))
    document.add_paragraph('Results: ')
    document.add_paragraph(precio)
    document.add_paragraph()
    document.add_paragraph('FIRM 1: ')
    document.add_paragraph(prod1)
    document.add_paragraph(beneficio1)
    document.add_paragraph()
    document.add_paragraph('FIRM 2: ')
    document.add_paragraph(prod2)
    document.add_paragraph(beneficio2)

    # Save the Word document
    document.save('bertrand_output.docx')


class BertrandMasInfoScreen(MasInfoVentana):
    a = '''Joseph Louis François Bertrand (1822-1900), a French mathematician and economist of the 19th century, \
revised the Cournot duopoly model, indicating that in a duopoly situation, companies would end up competing on prices, \
shaping what is known as the Bertrand model. \


The Bertrand model is a competition model between two or more companies in which, unlike in previous models, the \
decision variable is the price. That is, companies decide at what price they will sell their product, and the quantity \
will be determined by the market demand for that price. \


Starting from equal marginal costs, there is only one Nash equilibrium in the model: price_1 = price_2 = Marginal Cost \
since whenever the price is higher than the marginal cost, there are incentives to deviate from the situation. \


When companies have different marginal costs, the one with the lowest costs can capture the entire market demand and \
must decide between setting its monopoly price (if it can do so by excluding the rival company) or a price slightly \ 
below the marginal cost of the other company.'''
