from kivy.lang import Builder
from kivy.properties import StringProperty
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.textinput import TextInput
from main import VentanaLayout, MasInfoVentana
import sympy as sp
from kivy.garden.matplotlib.backend_kivyagg import FigureCanvasKivyAgg
import matplotlib.pyplot as plt
import numpy as np
from bimatrix_games import MensajeDeError
from docx import Document
from docx.shared import Inches


Builder.load_file('cournot.kv')


class CournotVentana(VentanaLayout):

    # Variables que mostraremos en la pantalla
    prod_1 = StringProperty('')
    prod_2 = StringProperty('')
    beneficio1 = StringProperty('')
    beneficio2 = StringProperty('')
    precio = StringProperty('')
    empresa1 = StringProperty('')
    empresa2 = StringProperty('')
    freacc1 = StringProperty('')
    freacc2 = StringProperty('')

    def obtenermasinfo(self):
        self.manager.current = 'CournotMasInfoScreen'

    def vacia(self):
        self.ids.a.text = ''
        self.ids.b.text = ''
        self.ids.c.text = ''
        self.ids.e.text = ''
        self.ids.graficournot.clear_widgets()
        self.prod_1 = ''
        self.prod_2 = ''
        self.beneficio1 = ''
        self.beneficio2 = ''
        self.precio = ''
        self.empresa1 = ''
        self.empresa2 = ''
        self.ids.fr1.text = ''
        self.ids.fr2.text = ''

    def aleatorio(self):
        self.ids.a.text = str(np.random.randint(20, 150))
        self.ids.b.text = str(np.random.randint(1, 10))
        self.ids.c.text = str(np.random.randint(1, 10))
        self.ids.e.text = str(np.random.randint(1, 10))

    def todo_en_orden(self):
        a = self.ids.a.text
        b = self.ids.b.text
        c = self.ids.c.text
        e = self.ids.e.text
        lista = [a, b, c, e]

        if '.' in lista or '' in lista or '0' in lista:
            alerta = MensajeDeError('Introduzca parámetros válidos')
            alerta.open()

        else:
            return True

    def calcula(self):
        if not self.todo_en_orden():
            return

        # Almacenamos en variables el input del usuario
        a = float(self.ids.a.text)
        b = float(self.ids.b.text)
        c = float(self.ids.c.text)
        e = float(self.ids.e.text)

        # Comprobamos si los parámetros son compatibles con una solución razonable del modelo
        if a < (2*c-e) or a < (2*e-c):
            alerta = MensajeDeError('Introduzca parámetro válidos')
            alerta.open()
            return

        # Definimos las variables del problema (cantidad producida por la empresa 1)
        x = sp.symbols('x')

        # Obtenemos las funciones de reacción de cada empresa, en función de x1 para poder graficarlo
        f1 = ((a - c)/b) - (2*x)
        f2 = ((a - e)/(2*b)) - (x/2)

        # Obtenemos las funciones de reacción como strings a mostrar
        self.freacc1 = f'{a - c}/{2*b}'
        self.freacc2 = f'{a - e}/{2*b}'

        # Obtenemos la producción óptima de cada empresa y pasamos el valor a la StringProperty a mostrar
        x1_sol = (a - (2*c) + e)/(3*b)
        x2_sol = (a - (2*e) + c)/(3*b)
        self.prod_1 = 'Producción óptima: ' + str(np.around(x1_sol, 3))
        self.prod_2 = 'Producción óptima: ' + str(np.around(x2_sol, 3))

        # Calculamos el beneficio de cada empresa y el precio final
        p = a - b*(x1_sol + x2_sol)
        self.beneficio1 = 'Beneficio: ' + str(np.around((p-c)*x1_sol, 3))
        self.beneficio2 = 'Beneficio: ' + str(np.around((p-e)*x2_sol, 3))
        self.precio = 'Precio en el mercado: ' + str(np.around(p, 3))
        self.empresa1 = 'EMPRESA 1'
        self.empresa2 = 'EMPRESA 2'

        # GRÁFICO
        # Convert the sp equation to a lambda function
        fl1 = sp.lambdify(x, f1)
        fl2 = sp.lambdify(x, f2)

        # Calculamos el mayor valor que pueden tomar x e y en cada función de reacción
        x1_max_f1 = (a-c)/(2*b)
        x1_max_f2 = (a-e)/b
        x2_max_f1 = (a-c)/b
        x2_max_f2 = (a-e)/(2*b)

        # escogemos un entero un 10% superior al máximo para graficarlo bien
        x1_max = int(np.ceil(1.1*np.maximum(x1_max_f1, x1_max_f2)))
        x2_max = int(np.ceil(1.1*np.maximum(x2_max_f1, x2_max_f2)))

        # Generamos los valores de x
        x_vals = np.linspace(0, x1_max, x1_max)

        # Calculamos los valores y asociados a cada valor de x mediante las funciones de reacción
        y1 = fl1(x_vals)
        y2 = fl2(x_vals)

        # Dibujamos el gráfico
        fig = plt.figure()
        ax = fig.add_subplot(111)

        # Dibujamos las ecuaciones
        ax.plot(x_vals, y1, label='f1')
        ax.plot(x_vals, y2, label='f2')

        # Dibujamos el punto de equilibrio
        ax.plot(x1_sol, x2_sol, marker='o', markersize=5, c='black', alpha=1)
        ax.annotate(f'Eq ({x1_sol:.3f}, {x2_sol:.3f})', xy=(x1_sol, x2_sol),
                    xycoords='data', xytext=(5, 5), textcoords='offset points')

        # Editamos las características del gráfico
        plt.title('Funciones de reacción')
        plt.xlabel('x$_{1}$')
        plt.ylabel('x$_{2}$')

        ax.set_xlim(0, x1_max)
        ax.set_ylim(0, x2_max)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.legend()

        # Líneas de puntos entre los ejes y el punto de equilibrio
        ax.plot([x1_sol, x1_sol], [0, x2_sol], ':', c='black')
        ax.plot([0, x1_sol], [x2_sol, x2_sol], ':', c='black')

        # Cambiamos el grosor de los bordes
        for spine in ['bottom', 'left']:
            ax.spines[spine].set_linewidth(2)

        # Creamos el widget FigureCanvasKivyAgg (vaciamos "graficournot" para eliminar el gráfico anterior)
        canvas = FigureCanvasKivyAgg(figure=fig)
        self.ids.graficournot.clear_widgets()
        self.ids.graficournot.add_widget(canvas)

        # Generamos el output en Word si el usuario lo requiere
        try:
            if self.ids.a_word.active:
                fdemanda = f'p(X) = {a} - {b}X ,    X = x\u2081 + x\u2082'
                ct1 = f'CT(x\u2081) = {c}x\u2081'
                ct2 = f'CT(x\u2082) = {e}x\u2082'

                fig.savefig('cournot_graph.png')
                crear_cournot_word(fdemanda, ct1, ct2, self.freacc1, self.freacc2, self.precio, self.prod_1,
                                   self.beneficio1, self.prod_2, self.beneficio2)
                alerta = MensajeDeError("Se ha creado el documento de Word 'cournot_output.docx' exitosamente")
                alerta.title = ''
                alerta.open()
        except PermissionError:
            alerta = MensajeDeError("Cierra los archivos con nombre 'cournot_output.docx' o 'cournot_graph.png'")
            alerta.open()


# clase para el estilo de cuadros de input de Cournot (5 caracteres, float positivo)
class CournotInput(TextInput):
    def insert_text(self, substring, from_undo=False):
        if len(self.text) >= 5:
            return
        try:
            if float(substring) < 0:
                return
        except ValueError:
            if substring == "." and "." in self.text:
                return
            elif substring != ".":
                return
        super().insert_text(substring, from_undo=from_undo)


class GrafiCournot(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)


# Método para exportar los resultados a Word
def crear_cournot_word(fdemanda, ct1, ct2, freacc1, freacc2, precio, prod1, beneficio1, prod2, beneficio2):
    # Create a new Word document
    document = Document()

    # Añadimos nuestros datos al documento
    document.add_paragraph('MODELO DE COURNOT')
    document.add_paragraph('Función de demanda del mercado: ' + fdemanda)
    document.add_paragraph('Costes totales de la empresa 1: ' + ct1)
    document.add_paragraph('Costes totales de la empresa 2: ' + ct2)
    document.add_paragraph()
    document.add_paragraph('Funciones de reacción: ')
    document.add_paragraph(f'Empresa 1: x\u2081(x\u2082) = {freacc1} - x\u2082/2 ')
    document.add_paragraph(f'Empresa 2: x\u2082(x\u2081) = {freacc2} - x\u2081/2 ')
    document.add_paragraph()
    document.add_picture('cournot_graph.png', width=Inches(6))
    document.add_paragraph('Resultados: ')
    document.add_paragraph(precio)
    document.add_paragraph()
    document.add_paragraph('EMPRESA 1: ')
    document.add_paragraph(prod1)
    document.add_paragraph(beneficio1)
    document.add_paragraph()
    document.add_paragraph('EMPRESA 2: ')
    document.add_paragraph(prod2)
    document.add_paragraph(beneficio2)

    # Save the Word document
    document.save('cournot_output.docx')


class CournotMasInfoScreen(MasInfoVentana):
    pass
