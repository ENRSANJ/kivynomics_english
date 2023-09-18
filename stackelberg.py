import numpy as np
from docx import Document
from docx.shared import Inches
from kivy.garden.matplotlib import FigureCanvasKivyAgg
from kivy.lang import Builder
import sympy as sp
from matplotlib import pyplot as plt
from bimatrix_games import MensajeDeError
from cournot import CournotVentana
from main import MasInfoVentana

Builder.load_file('stackelberg.kv')


# Heredamos la clase CournotVentana pues los modelos son muy parecidos y solo nos hará falta sobreescribir algún método
class StackelbergVentana(CournotVentana):

    # Sobreescribimos el método obtenermasinfo para abrir una ventana de más info distinta
    def obtenermasinfo(self):
        self.manager.current = 'StackelbergMasInfoScreen'

    # Sobreescribimos también el método "calcula" pues el equilibrio hallado es diferente
    def calcula(self):
        if not self.todo_en_orden():
            return

        # Almacenamos en variables el input del usuario
        a = float(self.ids.a.text)
        b = float(self.ids.b.text)
        c = float(self.ids.c.text)
        e = float(self.ids.e.text)

        # Comprobamos si los parámetros son compatibles con una solución razonable del modelo
        if a < (2 * c - e) or a < (3 * e - 2 * c):
            alerta = MensajeDeError('Enter valid parameters')
            alerta.open()
            return

        # Definimos las variables del problema (cantidad producida por la empresa líder, la 1)
        x = sp.symbols('x')

        # Obtenemos la función de reacción de la empresa 2
        f2 = ((a - e) / (2 * b)) - (x / 2)

        # Obtenemos la función de reacción como strings a mostrar
        self.freacc2 = f'{np.around(a - e, 3)}/{np.around(2*b, 3)}'

        # Obtenemos la producción óptima de cada empresa y pasamos el valor a la StringProperty a mostrar
        x1_sol = (a + e - 2 * c) / (2 * b)
        x2_sol = (a - 3 * e + 2 * c) / (4 * b)
        self.prod_1 = 'Optimal production: ' + str(np.around(x1_sol, 3))
        self.prod_2 = 'Optimal production: ' + str(np.around(x2_sol, 3))

        # Calculamos el beneficio de cada empresa y el precio final
        p = a - b * (x1_sol + x2_sol)
        self.beneficio1 = 'Profit: ' + str(np.around((p - c) * x1_sol, 3))
        self.beneficio2 = 'Profit: ' + str(np.around((p - e) * x2_sol, 3))
        self.precio = 'Market price: ' + str(np.around(p, 3))
        self.empresa1 = 'LEADER'
        self.empresa2 = 'FOLLOWER'

        # GRÁFICO
        # Convert the sympy equation to a lambda function
        fl2 = sp.lambdify(x, f2)

        # Calculamos el mayor valor que pueden tomar x e y
        x1_max_f2 = (a - e) / b
        x2_max_f2 = (a - e) / (2 * b)

        # escogemos un entero un 10% superior al máximo para graficarlo bien
        x1_max = int(np.ceil(1.1 * x1_max_f2 + 1))
        x2_max = int(np.ceil(1.1 * x2_max_f2 + 1))

        # Generamos los valores de x
        x_vals = np.linspace(0, x1_max, x1_max)

        # Calculamos los valores y asociados a cada valor de x mediante las funciones de reacción
        y2 = fl2(x_vals)

        # Dibujamos el gráfico
        fig = plt.figure()
        ax = fig.add_subplot(111)

        # Dibujamos las ecuaciones
        ax.plot(x_vals, y2, label='f2', color="orange")

        # Dibujamos el punto de equilibrio
        ax.plot(x1_sol, x2_sol, marker='o', markersize=5, c='black', alpha=1)
        ax.annotate(f'Eq ({x1_sol:.3f}, {x2_sol:.3f})', xy=(x1_sol, x2_sol),
                    xycoords='data', xytext=(5, 5), textcoords='offset points')

        # Editamos las características del gráfico
        plt.title('Reaction functions')
        plt.xlabel('x$_{1}$')
        plt.ylabel('x$_{2}$')

        ax.set_xlim(0, x1_max)
        ax.set_ylim(0, x2_max)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.legend()

        # Líneas de puntos entre los ejes y el punto de equilibrio
        ax.plot([x1_sol, x1_sol], [0, x2_sol], ':', c='black')
        ax.plot([0, x1_sol], [x2_sol, x2_sol], ':', c='black')

        # Cambiamos el grosor de los bordes
        for spine in ['bottom', 'left']:
            ax.spines[spine].set_linewidth(2)

        # Creamos el widget FigureCanvasKivyAgg (vaciamos "graficournot" para eliminar el gráfico anterior)
        canvas = FigureCanvasKivyAgg(figure=fig)
        plt.close("all")
        self.ids.graficournot.clear_widgets()
        self.ids.graficournot.add_widget(canvas)

        # Generamos el output en Word si el usuario lo requiere
        try:
            if self.ids.a_word.active:
                fdemanda = f'p(X) = {a} - {b}X ,    X = x\u2081 + x\u2082'
                ct1 = f'CT(x\u2081) = {c}x\u2081'
                ct2 = f'CT(x\u2082) = {e}x\u2082'

                fig.savefig('stackelberg_graph.png')
                crear_stackelberg_word(fdemanda, ct1, ct2, self.freacc2, self.precio, self.prod_1,
                                       self.beneficio1, self.prod_2, self.beneficio2)
                alerta = MensajeDeError("The Word document 'stackelberg_output.docx' has been succesfully created")
                alerta.title = ''
                alerta.open()
        except PermissionError:
            alerta = MensajeDeError(
                "Close all files named 'stackelberg_output.docx' or 'stackelberg_graph.png'")
            alerta.open()


# Método para exportar los resultados a Word
def crear_stackelberg_word(fdemanda, ct1, ct2, freacc2, precio, prod1, beneficio1, prod2, beneficio2):
    # Create a new Word document
    document = Document()

    # Añadimos nuestros datos al documento
    document.add_paragraph('STACKELBERG MODEL')
    document.add_paragraph('Market demand function: ' + fdemanda)
    document.add_paragraph('Total costs of firm 1: ' + ct1)
    document.add_paragraph('Total costs of firm 2: ' + ct2)
    document.add_paragraph()
    document.add_paragraph('Reaction functions: ')
    document.add_paragraph(f'Firm 2: x\u2082(x\u2081) = {freacc2} - x\u2081/2 ')
    document.add_paragraph()
    document.add_picture('stackelberg_graph.png', width=Inches(6))
    document.add_paragraph('Results: ')
    document.add_paragraph(precio)
    document.add_paragraph()
    document.add_paragraph('FIRM 1: ')
    document.add_paragraph(prod1)
    document.add_paragraph(beneficio1)
    document.add_paragraph()
    document.add_paragraph('FIRM 2: ')
    document.add_paragraph(prod2)
    document.add_paragraph(beneficio2)

    # Save the Word document
    document.save('stackelberg_output.docx')


class StackelbergMasInfoScreen(MasInfoVentana):
    a = '''Heinrich Freiherr von Stackelberg (1905-1946), a German economist, introduced his famous duopoly model in \
his work "Market Structure and Equilibrium," published in 1934. \

The Stackelberg duopoly, like the Cournot model, is also an economic model in which two firms compete by deciding on \
their production.

The firms determine the quantity produced sequentially; one firm (the leader) determines its production in the first \
period. This decision is irreversible and cannot be changed in the second period. Taking into account the leader's \
production, the other firm (the follower) determines its optimal production in the second period.

The model is, therefore, a dynamic game with perfect information: a player makes a decision after knowing the decision \
made by the other player.'''
