import numpy as np
import pygambit as pg
from kivy.animation import Animation
from kivy.lang import Builder
from kivy.uix.behaviors import DragBehavior
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.gridlayout import GridLayout
from kivy.uix.popup import Popup
from kivy.uix.scrollview import ScrollView
from kivy.uix.textinput import TextInput
from kivymd.theming import ThemableBehavior
from kivymd.uix.behaviors import HoverBehavior
from main import VentanaLayout, WrappedLabel, WrappedLabel2, MasInfoVentana
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


Builder.load_file('bimatrix_games.kv')


# Representaremos 3 decimales en cada output de la aplicación,,,,,,,,,,,,,,,EDITABLE EN OPCIONES?????
np.set_printoptions(formatter={'float': lambda x: "{0:0.3f}".format(x)})

# Filas y columnas a mostrar por defecto
rows = 2
cols = 2


# Mensaje de error por defecto
class MensajeDeError(Popup):
    def __init__(self, mensaje, **kwargs):
        super().__init__(**kwargs)

        self.title = '¡ERROR!'
        self.title_size = self.height * .2
        self.mensaje = mensaje
        self.size_hint = (.35, .35)
        msg = WrappedLabel(text=mensaje)
        msg.bind(height=lambda *args: setattr(msg, 'font_size', msg.height * .15))
        self.add_widget(msg)


# Mensaje de error específico, para errores en input de bimatrix_games
class MensajeDeErrorNash(MensajeDeError):
    def __init__(self, **kwargs):
        a = 'You must fill in all the positions of each payoff matrix with numerical values'
        super().__init__(mensaje=a, **kwargs)


# Popup con los equilibrios de Nash del juego
class PopupNash(DragBehavior, Popup):
    # to allow dragbehavior
    def _align_center(self, *args):
        pass


# Ventana principal de Juegos Bimatriciales
class BimatrixVentana(VentanaLayout):
    altura1 = .8
    espaciado1 = .105

    # Cálculo de los equilibrios de Nash
    def calculanash(self):
        lista1 = []
        lista2 = []

        # Recogemos los datos introducidos por el usuario
        try:
            for hijo in reversed(self.ids.matriz1.children):
                for nieto in reversed(hijo.children):
                    lista1.append(float(nieto.text))

            for hija in reversed(self.ids.matriz2.children):
                for nieta in reversed(hija.children):
                    lista2.append(float(nieta.text))

        # Los almacenamos en matrices de pagos, con el formato adecuado para pygambit
            payoff1 = [pg.Rational(num) for num in lista1]
            pagos1 = np.array(payoff1)
            matrizpagos1 = pagos1.reshape(rows, cols)

            payoff2 = [pg.Rational(num) for num in lista2]
            pagos2 = np.array(payoff2)
            matrizpagos2 = pagos2.reshape(rows, cols)
            juego = pg.Game.from_arrays(matrizpagos1, matrizpagos2)

            # Popup con Scrollview para mostrar el output
            scroll = ScrollView()
            solpopupnash = PopupNash(content=scroll)
            grid = GridLayout(cols=1, size_hint=(1, None), padding='.5dp')
            scroll.add_widget(grid)
            grid.bind(minimum_height=grid.setter('height'))

            # Para el output en Word
            listatotal2 = []

            # Mostramos como output los equilibrios de Nash, con un formato accesible
            for eq in pg.nash.enummixed_solve(juego, rational=False):
                strats = np.array(eq)
                eq_payoff = np.array2string(np.array([eq.payoff(player) for player in range(2)]), separator=', ')

                a = str('EE: [' + np.array2string(strats[:rows], separator=', ').replace('[', '(').replace(']', ')') +
                        ', ' + np.array2string(strats[cols:], separator=', ').replace('[', '(').replace(']', ')') +
                        ']  EP: ' + str(eq_payoff))

                c = WrappedLabel2(text=a, color=(0, 0, 0, 1), font_size=self.height * 0.022, size_hint=(1, None))
                grid.add_widget(c)
                listatotal2.append(a)

            solpopupnash.open()

            # Generamos el output en Word si el usuario lo requiere
            try:
                if self.ids.a_word.active:
                    listatotal = []
                    for i, j in zip(lista1, lista2):
                        listatotal.append((i, j))
                    crear_bimatrix_word(listatotal, rows, cols, listatotal2)
                    alerta = MensajeDeError("The Word document 'bimatrix_output.docx' has been succesfully created")
                    alerta.title = ''
                    alerta.open()
            except PermissionError:
                alerta = MensajeDeError("Close the  Word file named 'bimatrix_output.docx'")
                alerta.open()

        except ValueError:
            alertanash = MensajeDeErrorNash()
            alertanash.open()

    def sumanula(self):
        lista = []

        try:
            for hijo in self.ids.matriz1.children:
                for nieto in hijo.children:
                    lista.append(float(nieto.text))
            lista2 = np.array_split(lista, rows)
            for hija, sublista in zip(self.ids.matriz2.children, lista2):
                for nieta, a in zip(hija.children, sublista):
                    nieta.text = str(-a)

        except ValueError:
            alertanash = MensajeDeErrorNash()
            alertanash.open()

    def simetrico(self):
        lista = []
        if rows == cols:
            try:
                for hijo in self.ids.matriz1.children:
                    for nieto in hijo.children:
                        lista.append(float(nieto.text))
                lista2 = np.array(lista).reshape(rows, cols).transpose()
                for hija, sublista in zip(self.ids.matriz2.children, lista2):
                    for nieta, a in zip(hija.children, sublista):
                        nieta.text = str(a)
            except ValueError:
                alertanash = MensajeDeErrorNash()
                alertanash.open()
        else:
            matricesnosimetricas = MensajeDeError('Payment matrices must be square for a symmetrical game')
            matricesnosimetricas.open()

    def aleatorio(self):
        for hijo in self.ids.matriz1.children:
            for nieto in hijo.children:
                nieto.text = str(np.random.randint(-10, 10))

        for hija in self.ids.matriz2.children:
            for nieta in hija.children:
                nieta.text = str(np.random.randint(-10, 10))

    def vaciar(self):
        for hijo in self.ids.matriz1.children:
            for nieto in hijo.children:
                nieto.text = ''

        for hija in self.ids.matriz2.children:
            for nieta in hija.children:
                nieta.text = ''

    def obtenermasinfo(self):
        self.manager.current = 'BimatrixMasInfoScreen'


# BoxLayout para la selección del número de filas y columnas
class SeleccionaRango(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

    def confirmarango(self):
        global rows, cols
        try:

            if int(self.ids.numrows.text) in range(2, 7) and int(self.ids.numcols.text) in range(2, 7):
                rows = int(self.ids.numrows.text)
                cols = int(self.ids.numcols.text)
                Caja.cambiarango(self.parent.ids.matriz1)
                Caja.cambiarango(self.parent.ids.matriz2)
            else:
                errorrango = MensajeDeError('Enter a number of rows and columns between [2 and 6]')
                errorrango.open()

        except ValueError:
            errorrango = MensajeDeError('Enter a number of rows and columns between [2 and 6]')
            errorrango.open()

        self.ids.numrows.text = ''
        self.ids.numcols.text = ''


# Clase para el botón de "Ok" de SeleccionaRango
class OkButton(Button, ThemableBehavior, HoverBehavior):

    def on_enter(self, *args):
        Animation(size_hint=(1, 1), d=0.1).start(self)

    def on_leave(self, *args):
        Animation(size_hint=(1, .9), d=0.1).start(self)


# BoxLayout que contendrá cada una de las matrices en pantalla
class Caja(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'
        self.cambiarango()
        self.spacing = '10dp'

    def cambiarango(self):
        self.clear_widgets()
        for x in range(rows):
            a = Fila()
            self.add_widget(a)

        for x in self.children:
            for y in range(cols):
                x.add_widget(Cuadradito())


# Fila de las matrices (contendrá a los "Cuadraditos")
class Fila(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.spacing = '20dp'


# TextInput para cada elemento de las matrices
class Cuadradito(TextInput):
    pass


# Textinput para la selección del número de filas y columnas de las matrices de pagos
class SeleccionaInput(TextInput):
    pass


class BimatrixMasInfoScreen(MasInfoVentana):
    # Texto a mostrar. Escrito aquí porque Kivy lang no lee los triple-quoted strings de la misma forma
    a = '''In the 1950s, John Forbes Nash (1928-2015), an American mathematician and Nobel laureate in Economics in \
1994, introduced very relevant concepts in the field of game theory (Nash equilibrium, Nash bargaining). \
Additionally, he demonstrated the existence of at least one Nash equilibrium in every finite game.\


The output provided by the application consists of:\


EE: (Extreme Equilibria) probabilities assigned by each player to their strategies in the Nash equilibrium\


EP: (Expected Payoff) Expected payoff for each player in the equilibrium'''

    b = '''In this example, we have 2 equilibria as a result:\
    

1st: It is a pure strategy equilibrium where the players play their second and first strategies respectively, \
obtaining payoffs of 6 and 7.


2nd: It is a mixed strategy equilibrium, where each player plays their moves 1 and 2 with different probabilities.'''


# Método para exportar los resultados a Word
def crear_bimatrix_word(datos, filas, columnas, equilibrios):
    # Create a new Word document
    document = Document()

    # Add some text to the document
    document.add_paragraph('BIMATRIX GAMES')

    # Create a table with the desired shape

    table = document.add_table(rows=filas, cols=columnas)

    # Fill in the cells of the table with the values from the data list
    for i in range(filas):
        row_cells = table.rows[i].cells
        for j in range(columnas):
            value = datos[i * cols + j]
            row_cells[j].text = str(value)

    # Set table properties
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = True

    # Set cell properties
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1)

    # Añadimos los equilibrios de Nash
    document.add_paragraph('')
    document.add_paragraph('Nash equilibria: ')
    for i in equilibrios:
        document.add_paragraph(i)

    # Save the Word document
    document.save('bimatrix_output.docx')
